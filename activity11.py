import os
import math
from datetime import datetime
from docx import Document
import wikipedia
import simpleeval

from langchain_core.tools import tool
from langchain_openai import ChatOpenAI
from langgraph.prebuilt import create_react_agent

# =====================================================================
# 1. LLM SETUP
# =====================================================================
llm = ChatOpenAI(
    base_url="http://localhost:1234/v1",
    api_key="lm-studio",
    model="google/gemma-4-e4b",
    temperature=0.2
)

# =====================================================================
# 2. FILE CREATION HELPER
# =====================================================================
def create_word_doc(content: str, filename: str) -> str:
    """Helper function to format and save a Word document."""
    folder = "generated_docs"
    os.makedirs(folder, exist_ok=True)
    filepath = os.path.join(folder, filename)

    doc = Document()
    doc.add_heading("AI Generated Document", level=0)

    for line in content.split("\n"):
        if line.strip():
            doc.add_paragraph(line)

    doc.save(filepath)
    return f"Success! Document saved at: {filepath}"

# =====================================================================
# 3. DEFINE THE AGENT TOOLS
# =====================================================================

@tool
def calculator_tool(expression: str) -> str:
    """Solve math/arithmetic problems like addition, subtraction, multiplication, division, and square roots."""
    try:
        evaluator = simpleeval.SimpleEval()
        evaluator.functions = {
            "sqrt": math.sqrt,
            "pow":  math.pow,
            "abs":  abs,
        }
        evaluator.names = {"pi": math.pi, "e": math.e}
        result = evaluator.eval(expression)
        return f"Result: {result}"
    except simpleeval.InvalidExpression as e:
        return f"Invalid expression: {str(e)}"
    except Exception as e:
        return f"Could not evaluate expression. Error: {str(e)}"


@tool
def wikipedia_tool(query: str) -> str:
    """Search Wikipedia and return a short summary of any general knowledge topic."""
    wikipedia.set_lang("en")
    try:
        return wikipedia.summary(query, sentences=4, auto_suggest=True)
    except wikipedia.exceptions.DisambiguationError as e:
        options = ", ".join(e.options[:5])
        return f"'{query}' is ambiguous. Did you mean one of these? {options}. Please retry with a more specific term."
    except wikipedia.exceptions.PageError:
        return f"No Wikipedia page found for '{query}'. Try rephrasing the query."
    except Exception as e:
        return f"Wikipedia lookup failed for '{query}'. Error: {str(e)}"


@tool
def word_document_tool(query: str) -> str:
    """Creates a stylized Word document (.docx) based on a general topic or raw text requested by the user."""
    prompt = f"Write structured document content about: {query}. Include sections and bullet points."
    content = llm.invoke(prompt).content

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"doc_{timestamp}.docx"

    return create_word_doc(content, filename)


@tool
def wiki_to_document_tool(query: str) -> str:
    """Searches Wikipedia for a topic and generates a saved Word document from it.
    Use this when the user specifically asks to make/create a document out of a Wikipedia topic."""
    wikipedia.set_lang("en")
    try:
        summary = wikipedia.summary(query, sentences=10, auto_suggest=True)
    except wikipedia.exceptions.DisambiguationError as e:
        options = ", ".join(e.options[:5])
        return f"Could not build a document — '{query}' is ambiguous. Try one of these instead: {options}."
    except wikipedia.exceptions.PageError:
        return f"No Wikipedia page found for '{query}'. Try rephrasing the topic."
    except Exception as e:
        return f"Could not pull Wikipedia data for '{query}'. Error: {str(e)}"

    sentences = summary.split('. ')
    intro_paragraph = ". ".join(sentences[:3]) + "."

    bullet_points = ""
    for s in sentences[3:]:
        if s.strip():
            bullet_points += f"• {s.strip()}.\n"

    content = f"""RESEARCH DOSSIER: {query.upper()}
Generated on: {datetime.now().strftime('%B %d, %Y')}

1. Executive Overview
{intro_paragraph}

2. Key Facts & Data Points
{bullet_points}
3. Archive Metadata
- Source: Wikipedia API Extraction
- System: Integrated Agent Document Engine
"""

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"wiki_doc_{timestamp}.docx"

    return create_word_doc(content, filename)


# Collect all tools
tools = [calculator_tool, wikipedia_tool, word_document_tool, wiki_to_document_tool]

# =====================================================================
# 4. INITIALIZE THE AGENT
# =====================================================================
system_instruction = """You are an intelligent terminal assistant.
You have access to tools for Math, Wikipedia lookups, text generation, and Wikipedia document archives.

DECISION RULES:
1. If the user wants a simple summary of a topic -> Use 'wikipedia_tool'.
2. If the user wants to compile a topic into a Word Document from Wikipedia data -> Use 'wiki_to_document_tool'.
3. If the user wants to create an arbitrary text document from scratch -> Use 'word_document_tool'.
4. If the user asks for math calculations -> Use 'calculator_tool'.

If a Wikipedia tool returns an ambiguity or page error, inform the user and suggest the alternatives provided.
Always provide a final answer stating clearly what action you took or what file was created.
"""

agent_executor = create_react_agent(
    model=llm,
    tools=tools,
    prompt=system_instruction
)

# =====================================================================
# 5. CHAT LOOP RUNNER
# =====================================================================
if __name__ == "__main__":
    print("==================================================")
    print("AI Multi-Tool Agent Active! Type 'exit' to close.")
    print("==================================================")

    while True:
        user_input = input("\nYou: ").strip()
        if user_input.lower() == "exit":
            print("Closing workspace. Goodbye!")
            break

        if not user_input:
            continue

        print("\nThinking...")

        try:
            response = agent_executor.invoke({"messages": [("user", user_input)]})
            final_reply = response["messages"][-1].content
        except Exception as e:
            final_reply = f"Agent encountered an error: {str(e)}"

        print("\n" + "-" * 50)
        print(f"AI: {final_reply}")
        print("-" * 50)